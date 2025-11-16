"""
MarkdownをWord文書（.docx）に変換するスクリプト
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def parse_markdown_to_docx(md_file: str, docx_file: str):
    """MarkdownファイルをWord文書に変換"""

    # Word文書を作成
    doc = Document()

    # スタイル設定
    setup_styles(doc)

    # Markdownファイルを読み込み
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # コードブロックの処理
        if line.startswith('```'):
            if in_code_block:
                # コードブロック終了
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # コードブロック開始
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # テーブルの処理
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # 次の行がテーブルでない場合、テーブル終了
            if i >= len(lines) or '|' not in lines[i]:
                add_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue

        # 見出しの処理
        if line.startswith('#'):
            add_heading(doc, line)
        # 水平線の処理
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        # 箇条書きの処理
        elif line.startswith('- ') or line.startswith('* '):
            add_bullet(doc, line)
        # 番号付きリストの処理
        elif re.match(r'^\d+\. ', line):
            add_numbered_list(doc, line)
        # 空行
        elif not line.strip():
            if i > 0 and lines[i-1].strip():  # 連続する空行は1つだけ
                doc.add_paragraph()
        # 通常のテキスト
        else:
            add_paragraph(doc, line)

        i += 1

    # 保存
    doc.save(docx_file)
    print(f"Word文書を作成しました: {docx_file}")


def setup_styles(doc):
    """スタイルを設定"""
    styles = doc.styles

    # 見出しスタイルの調整
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Yu Gothic'
            style.font.size = Pt(20 - level * 2)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 51, 102)


def add_heading(doc, line):
    """見出しを追加"""
    level = 0
    while line[level] == '#':
        level += 1

    text = line[level:].strip()

    # 見出しレベルを制限（最大3）
    level = min(level, 3)

    doc.add_heading(text, level=level)


def add_paragraph(doc, line):
    """段落を追加（太字、コード、リンクの処理）"""
    p = doc.add_paragraph()

    # インラインマークダウンを処理
    parts = parse_inline_markdown(line)

    for part_type, text in parts:
        run = p.add_run(text)

        if part_type == 'bold':
            run.bold = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        elif part_type == 'italic':
            run.italic = True


def parse_inline_markdown(line):
    """インラインマークダウンをパース"""
    parts = []
    current_text = ""
    i = 0

    while i < len(line):
        # 太字 **text**
        if line[i:i+2] == '**':
            if current_text:
                parts.append(('normal', current_text))
                current_text = ""

            # 終了の**を探す
            end = line.find('**', i+2)
            if end != -1:
                parts.append(('bold', line[i+2:end]))
                i = end + 2
                continue

        # コード `text`
        elif line[i] == '`':
            if current_text:
                parts.append(('normal', current_text))
                current_text = ""

            # 終了の`を探す
            end = line.find('`', i+1)
            if end != -1:
                parts.append(('code', line[i+1:end]))
                i = end + 1
                continue

        current_text += line[i]
        i += 1

    if current_text:
        parts.append(('normal', current_text))

    return parts


def add_bullet(doc, line):
    """箇条書きを追加"""
    text = line[2:].strip()  # '- ' or '* ' を削除
    parts = parse_inline_markdown(text)

    p = doc.add_paragraph(style='List Bullet')
    for part_type, text in parts:
        run = p.add_run(text)
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)


def add_numbered_list(doc, line):
    """番号付きリストを追加"""
    match = re.match(r'^\d+\. (.+)$', line)
    if match:
        text = match.group(1)
        doc.add_paragraph(text, style='List Number')


def add_code_block(doc, code_lines):
    """コードブロックを追加"""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph(code_text)
    p.style = 'No Spacing'

    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)

    # 背景色（グレー）を設定
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)


def add_table(doc, table_lines):
    """テーブルを追加"""
    # テーブル行をパース
    rows = []
    for line in table_lines:
        # セパレーター行をスキップ
        if re.match(r'^\|[\s\-:]+\|$', line):
            continue

        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return

    # テーブルを作成
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # データを設定
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data

            # ヘッダー行は太字
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def main():
    """メイン実行"""
    import sys

    # コマンドライン引数から入力ファイルを取得
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
        docx_file = md_file.replace('.md', '.docx')
    else:
        md_file = 'OVERTIME_CALCULATION_GUIDE.md'
        docx_file = 'OVERTIME_CALCULATION_GUIDE.docx'

    print(f"Markdown → Word変換を開始: {md_file}")
    parse_markdown_to_docx(md_file, docx_file)
    print("変換完了！")


if __name__ == '__main__':
    main()
